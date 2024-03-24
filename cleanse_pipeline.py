import re
import copy
import spacy
import pandas as pd
import language_tool_python
from pathlib import Path
from docx import Document
from docx.shared import Cm
from itertools import product
from functools import partial
from multiprocessing import Pool

# from generate_rels import *

# tool = language_tool_python.LanguageTool(
#     "en-US"
# )  # use a local server (automatically set up), language English

# Load the English language model
nlp = spacy.load("en_core_web_sm")


def starts_with_figure(row):
    pattern = r"^(Figure|Fig)\s*\d+[.:]"
    return re.match(pattern, row) is not None


def starts_with_table(row):
    pattern = r"^(Table|Tab)\s*\d+[.:]"
    return re.match(pattern, row) is not None


def is_finished(s):
    pattern = r"\.$"
    return bool(re.search(pattern, s))


def remove_consecutive_whitespaces(text):
    # Use regular expression to replace consecutive whitespaces with a single space
    cleaned_text = re.sub(r"\s+", " ", text)
    return cleaned_text


def remove_spaces_around_parentheses(text):
    # Remove spaces after opening parenthesis, bracket, or brace and before a word
    text = re.sub(r"(?<=[\(\[\{])\s+(?=\w)", "", text)
    # Remove spaces after a word and before closing parenthesis, bracket, or brace
    text = re.sub(r"(?<=\w)\s+(?=[\)\]\}])", "", text)
    return text


def select_text_smaller(document, font_size=8):
    # this is ok for we don't are not currently using the
    doc = copy.deepcopy(document)
    selected_text = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check if the font size is smaller than font_size points
            if run.font.size and run.font.size.pt <= font_size:
                run._element.getparent().remove(run._element)
                # selected_text.append(run.text)
                # if run.text != "":
                #     print(run.text)
    return doc


def select_text_greater(document, font_size=13):
    # this is ok for we don't are not currently using the
    doc = copy.deepcopy(document)
    selected_text = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check if the font size is smaller than font_size points
            if run.font.size and run.font.size.pt > font_size:
                run._element.getparent().remove(run._element)
                # selected_text.append(run.text)
                # if run.text != "":
                #     print(run.text)
    return doc


def merge_lowercase_with_hyphen(text):
    # Define a regex pattern to match lowercase strings containing hyphens
    pattern = r"\b[a-z]+-[a-z]+\b"

    # Find all matches in the text
    matches = re.findall(pattern, text)

    # Replace each match with a merged string
    for match in matches:
        # Merge lowercase strings with hyphen
        merged_string = match.replace("-", "")
        # Replace the match with the merged string in the original text
        text = text.replace(match, merged_string)

    return text


def delete_images_and_tables(document):

    # doc = document.copy()
    doc = copy.deepcopy(document)
    images_to_delete = []
    tables_to_delete = []

    for idx, paragraph in enumerate(doc.paragraphs):
        # Check for images
        for run in paragraph.runs:
            if (
                run._element.tag
                == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
            ):
                for pic in run._element.iter(
                    "{http://schemas.openxmlformats.org/drawingml/2006/picture}pic"
                ):
                    images_to_delete.append((idx, pic))

    # Remove images
    for idx, pic in images_to_delete:
        doc.paragraphs[idx].clear()

    # Identify tables to delete
    tables_to_delete = []
    for idx, table in enumerate(doc.tables):
        tables_to_delete.append(idx)

    # Remove tables
    for idx in reversed(tables_to_delete):
        doc._element.body.remove(doc.tables[idx]._element)

    return doc


def text_to_sentences(text):
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents]
    return sentences


def correct_morphology(text):
    matches = tool.check(text)
    # select the matches to be applid to the text
    for m in matches:
        if m.ruleIssueType == "misspelling":
            try:
                m.replacements = [m.replacements[1]]
            except IndexError:
                continue
        else:
            matches.remove(m)
    corrected_text = language_tool_python.utils.correct(text, matches)
    return corrected_text


def pipeline(document):

    doc = delete_images_and_tables(document)
    # doc = select_text_smaller(doc)
    # doc = select_text_greater(doc)
    doc = document

    attachments = []
    show = False
    data = []
    aux = []
    keywords = ["abstract", "introduction", "summary", "background", "a b s t r a c t"]
    # is_present = any(keyword in p.text.lower() for keyword in keywords)
    last_data = ""
    for p in doc.paragraphs:
        
        # if p.text.startswith("Similarly"):
        #     # continue

        if p.text == "":
            continue
        # print(p.text)
        #         print(show)

        # print("p.text 1", p.text)
        if starts_with_table(p.text) or starts_with_figure(p.text):
            attachments.append(p.text)
            continue
        
        # print("p.text 2")
        if p.text.upper() == "REFERENCES":
            show = False
            break
        # print("p.text 2")

        aux.append(p.text)

        if show is False:
            if any(keyword in p.text.lower() for keyword in keywords):
                show = True
                data.append(p.text)
                continue

        if p.text.isnumeric():
            continue

        # if p.text.upper() == p.text:
        #     # print(p.text)
        #     continue
        
        # print("p.text 2")

        if show:
            # print(p.text)
            # if is_finished(p.text):
            #     # print("is finished")
            #     last_data += " " + p.text
            #     # print(last_data)
            #     # continuous_text = merge_lowercase_with_hyphen(last_data)
            #     # morph_text = correct_morphology(continuous_text)
            #     # last_data = remove_consecutive_whitespaces(last_data).strip()
            #     last_data = last_data.strip()
            #     # clean_text = remove_spaces_around_parentheses(clean_text)
            #     # clean_text = clean_text.strip()
            #     # clean_text = re.sub(r"\n+", " ", clean_text)
            #     # clean_sentences = clean_text.strip().replace(". ", ".<|>").split("<|>")

            #     # clean_sentences = text_to_sentences(clean_text)

            #     # data.extend(clean_sentences)
            #     # print(last_data)
            #     # print()
            #     data.append(last_data)

            #     last_data = ""
            # else:
            # last_data = p.text.replace("-\n", "")
            data.append(p.text)
            continue

        # print("p.text 3", show)
        # print(p.text, show)
        # if show:
        #     data.append(p.text)
        # print(last_data)
    # print(attachments)
    # data.extend(attachments)
    if len(data) == 0:
        return aux
    else:
        return data


input_folder = Path("dataset_docx_ocr")

input_files = input_folder.glob("*.docx")

# input_files = [
#     Path("dataset_docx_ocr/c78dc997dcac438e41018967eda07b4b.docx"),
#     Path("dataset_docx_ocr/5dc298c31c827ab090520d8f613d945d.docx"),
#     Path("dataset_docx_ocr/84e018e77c6b5a5df94845286ce8e2e7.docx"),
#     Path("dataset_docx_ocr/043e94ea8c7f258bbb81ce6fe1f8bfcf.docx"),
#     Path("dataset_docx_ocr/4bde6692f1c62c246729f22cd3a0c09d.docx"),
#     Path("dataset_docx_ocr/212c7af70a5700b7db626d93dbc8a24f.docx")
#     Path("dataset_docx_ocr/1d9eecbac010038b00fd49fc723ac849.docx") # empty
# ]

# input_files = [Path("dataset_docx_ocr/0a337a62de92b3e2717d0f7a454dbe4c.docx")]
# input_files = [Path("docx_2_text_1/0ddcfd1cb232030765e364e5064ddcb9.txt")]
# # dataset = []
# text_rels_pairs = []

for f in input_files:
    # if Path(f"docx_2_text_1/{f.stem}.txt").exists():
    #     continue
    document = Document(f)
    data = pipeline(document)

    # for d in data:
    #     # rels = ner_extract(d)
    #     row = {
    #         "descriptor": f.stem,
    #         "text": d,
    #         # "ner": rels
    #     }
    #     text_rels_pairs.append(row)
    data = " ".join(data)
    with open(f"docx_2_txt_2/{f.stem}.txt", "w") as fd:
        fd.write(data)
#     # data = data.replace("-\n", "")
#     # data = remove_consecutive_whitespaces(data)
#     # print(data)

#     prompt = f"""given the following text:

# {data}

# correct the punctuation, spelling and typos. place the blocks of text that prevents the sentences to be semantically coherent, maintain the original sentence, to the end of the text after a `\n\n` pause.
# make sure that no content is lost nor added.

# is the output correct?
# """
#     # print(name, len(data))
#     # print(data)
#     print()
#     print()
#     with open(f"prompts/{f.stem}.txt", "w") as f:
#         f.write(prompt)
#     # f_name = f.stem
    # data = list(product([f_name], data))
    # dataset.extend(data)
    # print(pd.DataFrame(text_rels_pairs))

# df = pd.DataFrame(text_rels_pairs)
# print(df)
# df.to_csv("doc_clean_stripped_f_abstract.csv", index=None)


# def process_document(f):
#     document = Document(f)
#     data = pipeline(document)  # Assuming you have defined your pipeline function
#     f_name = f.stem
#     data = list(product([f_name], data))
#     return data

# # Define the number of processes
# num_processes = 8  # Adjust as needed

# # Create a Pool of processes
# with Pool(num_processes) as pool:
#     # Map the process_document function to each input file
#     results = pool.map(process_document, input_files)

# # Flatten the results list
# dataset = [item for sublist in results for item in sublist]

# # Create a DataFrame from the dataset
# df = pd.DataFrame(dataset)

# # Save the DataFrame to CSV
# df.to_csv("doc_clean_parallel_data.csv", index=None)


# chunked_folder = Path("chunked_4000")
# chunked_files = chunked_folder.glob("*.txt")

# for f in chunked_files:

#     with open(f, "r") as fd:
#         text = fd.read()

#     prompt = f"""given the text below:

# {text}

# correct the punctuation, spelling, orthography and typos.
# then, place the blocks of text that prevent the sentences to be semantically coherent, to the end of the text after a line break followed by`gobbledygook:` pause.
# the corrected text should contain no line breaks, no symbols that are not present in the original text.
# the corrected text is a standard text file, no markdown language should be present.
# get rid of any table or tabular data, e.g.: 2 (11.11) 16 (88.89) 5 (27.78) 2 (11.11) 11 (61.11) 2 (11.11) 16 (88.89)
# get rid of all references.
# get rid of all text that is detrimental for the contextual understanding of the text.
# make sure that no information is lot or added, but re-arranged.
#     """
#     with open(f"prompt_4000/{f.stem}.txt", "w") as fd:
#         fd.write(prompt)